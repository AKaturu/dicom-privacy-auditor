from __future__ import annotations

import json
from copy import deepcopy
from importlib.resources import files

import pytest
from pydicom.dataset import Dataset, FileDataset, FileMetaDataset
from pydicom.sequence import Sequence
from pydicom.uid import ExplicitVRLittleEndian, SecondaryCaptureImageStorage, generate_uid

from dicom_privacy_auditor.ps315 import (
    PolicySelection,
    ProfileOption,
    all_code_rules,
    all_rules,
    evaluate_pair,
    get_rule,
    resolve_rule,
    table_metadata,
)


def _write(path, dataset: Dataset) -> None:
    meta = FileMetaDataset()
    meta.MediaStorageSOPClassUID = SecondaryCaptureImageStorage
    meta.MediaStorageSOPInstanceUID = str(dataset.SOPInstanceUID)
    meta.TransferSyntaxUID = ExplicitVRLittleEndian
    file_dataset = FileDataset(str(path), {}, file_meta=meta, preamble=b"\0" * 128)
    file_dataset.update(dataset)
    file_dataset.save_as(path, enforce_file_format=True)


def test_user_generated_tables_are_loaded_from_configured_directory():
    metadata = table_metadata()
    assert metadata["edition"] == "test"
    assert metadata["generated_locally"] is True
    assert metadata["redistributed_by_project"] is False
    assert metadata["row_count"] == 8
    assert metadata["code_rule_count"] == 1
    assert len(all_rules()) == 8
    assert len(all_code_rules()) == 1


def test_complete_tables_are_not_bundled_in_package():
    data = files("dicom_privacy_auditor.data")
    assert not any(item.name.startswith("ps315_") and item.name.endswith(".json") for item in data.iterdir())


def test_missing_local_tables_produces_actionable_error(tmp_path, monkeypatch):
    from dicom_privacy_auditor.ps315.policy import (
        StandardsDataNotInstalledError,
        clear_caches,
        load_table,
    )

    monkeypatch.setenv("DICOM_PRIVACY_PS315_DATA_DIR", str(tmp_path))
    monkeypatch.setenv("DICOM_PRIVACY_PS315_EDITION", "missing")
    clear_caches()
    with pytest.raises(StandardsDataNotInstalledError, match="prepare-data"):
        load_table()
    monkeypatch.undo()
    clear_caches()


def test_option_overrides_are_preserved_without_inventing_precedence():
    rule = get_rule("0020000D")  # Study Instance UID
    assert rule is not None
    basic = resolve_rule(rule, PolicySelection())
    retained = resolve_rule(rule, PolicySelection(options=(ProfileOption.RETAIN_UIDS,)))
    assert "U" in basic.directives[0]
    assert retained.directives == ("K",)


def test_pair_evaluation_is_recursive_and_checks_private_attributes(tmp_path):
    source = Dataset()
    source.SOPClassUID = SecondaryCaptureImageStorage
    source.SOPInstanceUID = generate_uid()
    source.PatientName = "DOE^JANE"
    source.StudyInstanceUID = generate_uid()
    source.add_new((0x0043, 0x0010), "LO", "PRIVATE CREATOR")
    source.add_new((0x0043, 0x1029), "LO", "SECRET")
    item = Dataset()
    item.PatientID = "MRN123"
    source.RequestAttributesSequence = Sequence([item])

    candidate = deepcopy(source)
    candidate.PatientName = ""
    candidate.SOPInstanceUID = generate_uid()
    candidate.StudyInstanceUID = generate_uid()
    del candidate[(0x0043, 0x0010)]
    del candidate[(0x0043, 0x1029)]
    candidate.RequestAttributesSequence[0].PatientID = ""

    source_path = tmp_path / "source.dcm"
    candidate_path = tmp_path / "candidate.dcm"
    _write(source_path, source)
    _write(candidate_path, candidate)
    evaluation = evaluate_pair(source_path, candidate_path)
    nested = next(result for result in evaluation.results if result.path.endswith("PatientID"))
    assert nested.status == "pass"
    assert any(result.rule_name == "Private Attributes" for result in evaluation.results)
    assert all("source.dcm" not in value for value in (evaluation.source, evaluation.candidate))


def test_sr_code_table_is_applied(tmp_path):
    source = Dataset()
    source.SOPClassUID = SecondaryCaptureImageStorage
    source.SOPInstanceUID = generate_uid()
    item = Dataset()
    item.ValueType = "TEXT"
    code = Dataset()
    code.CodeValue = "121022"  # Accession Number, action X in E.1-2
    code.CodingSchemeDesignator = "DCM"
    code.CodeMeaning = "Accession Number"
    item.ConceptNameCodeSequence = Sequence([code])
    item.TextValue = "ACCESSION-123"
    source.ContentSequence = Sequence([item])
    candidate = deepcopy(source)
    del candidate.ContentSequence

    source_path = tmp_path / "sr-source.dcm"
    candidate_path = tmp_path / "sr-candidate.dcm"
    _write(source_path, source)
    _write(candidate_path, candidate)
    evaluation = evaluate_pair(source_path, candidate_path)
    code_results = [result for result in evaluation.results if result.tag == "CODE"]
    assert len(code_results) == 1
    assert code_results[0].status == "pass"
    assert code_results[0].observed == "X"


def test_empty_sequence_is_classified_as_zero_length():
    from dicom_privacy_auditor.ps315.evaluate import _observed

    source = Dataset()
    source.RequestAttributesSequence = Sequence([Dataset()])
    candidate = Dataset()
    candidate.RequestAttributesSequence = Sequence([])
    assert _observed(source[0x00400275], candidate[0x00400275]) == "Z"


def test_candidate_only_forbidden_private_attribute_is_flagged(tmp_path):
    source = Dataset()
    source.SOPClassUID = SecondaryCaptureImageStorage
    source.SOPInstanceUID = generate_uid()
    candidate = deepcopy(source)
    candidate.add_new((0x0043, 0x0010), "LO", "PRIVATE CREATOR")
    candidate.add_new((0x0043, 0x1029), "LO", "NEW SECRET")

    source_path = tmp_path / "source.dcm"
    candidate_path = tmp_path / "candidate.dcm"
    _write(source_path, source)
    _write(candidate_path, candidate)
    evaluation = evaluate_pair(source_path, candidate_path)
    introduced = [
        item
        for item in evaluation.results
        if item.path in {"(0043,0010)", "(0043,1029)"} or item.rule_name == "Private Attributes"
    ]
    assert any(item.status == "fail" and item.observed == "PRESENT" for item in introduced)


def test_file_meta_uid_mismatch_is_an_operational_failure():
    from dicom_privacy_auditor.ps315.evaluate import _file_meta_consistency_checks

    candidate = Dataset()
    candidate.SOPClassUID = SecondaryCaptureImageStorage
    candidate.SOPInstanceUID = generate_uid()
    candidate.file_meta = FileMetaDataset()
    candidate.file_meta.MediaStorageSOPClassUID = SecondaryCaptureImageStorage
    candidate.file_meta.MediaStorageSOPInstanceUID = generate_uid()
    checks = _file_meta_consistency_checks(candidate)
    check = next(item for item in checks if item.get("check") == "file_meta_sop_instance_uid_consistency")
    assert check["status"] == "fail"


def test_local_docx_generator_writes_user_cache_without_copying_source(tmp_path):
    from docx import Document

    from dicom_privacy_auditor.ps315.generate import ATTRIBUTE_COLUMNS, CODE_COLUMNS, generate_tables

    source = tmp_path / "official-user-copy.docx"
    document = Document()
    document.add_paragraph("DICOM PS3.15 testgen")
    attribute_table = document.add_table(rows=2, cols=len(ATTRIBUTE_COLUMNS))
    for index, value in enumerate(
        [
            "Attribute Name",
            "Tag",
            "Retd. (Y/N)",
            "In Std. Comp. IOD (Y/N)",
            "Basic Prof.",
            "Rtn. Safe Priv. Opt.",
            "Rtn. UIDs Opt.",
            "Rtn. Dev. Id. Opt.",
            "Rtn. Inst. Id. Opt.",
            "Rtn. Pat. Chars. Opt.",
            "Rtn. Long. Full Dates Opt.",
            "Rtn. Long. Modif. Dates Opt.",
            "Clean Desc. Opt.",
            "Clean Struct. Cont. Opt.",
            "Clean Graph. Opt.",
        ]
    ):
        attribute_table.rows[0].cells[index].text = value
    attribute_values = [
        "Synthetic rule",
        "(0010,0010)",
        "N",
        "Y",
        "Z",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
    ]
    for index, value in enumerate(attribute_values):
        attribute_table.rows[1].cells[index].text = value

    code_table = document.add_table(rows=2, cols=len(CODE_COLUMNS))
    for index, value in enumerate(
        [
            "Code Meaning",
            "Code Value",
            "Coding Scheme Designator",
            "Value Type",
            "Retd. (Y/N)",
            "In Std. TID (Y/N)",
            "Basic Prof.",
            "Rtn. UIDs Opt.",
            "Rtn. Dev. Id. Opt.",
            "Rtn. Inst. Id. Opt.",
            "Rtn. Pat. Chars. Opt.",
            "Rtn. Long. Full Dates Opt.",
            "Rtn. Long. Modif. Dates Opt.",
            "Clean Desc. Opt.",
        ]
    ):
        code_table.rows[0].cells[index].text = value
    code_values = ["Synthetic code", "1", "99TEST", "TEXT", "N", "Y", "X", "", "", "", "", "", "", ""]
    for index, value in enumerate(code_values):
        code_table.rows[1].cells[index].text = value
    document.save(source)

    output = tmp_path / "cache"
    generated = generate_tables(
        source,
        edition="testgen",
        output=output,
        minimum_attribute_rows=1,
        minimum_code_rows=1,
    )
    assert all(path.is_file() for path in generated)
    assert not (output / source.name).exists()
    payload = json.loads(generated[0].read_text(encoding="utf-8"))
    assert payload["generated_locally"] is True
    assert payload["redistributed_by_project"] is False
