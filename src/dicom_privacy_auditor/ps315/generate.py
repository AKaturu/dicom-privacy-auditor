"""Generate a local DICOM PS3.15 rule cache from an official Part 15 DOCX.

The DICOM Standard is not redistributed by this project. Users may point this
module at a locally obtained official document or ask it to download the
current official DOCX temporarily from the DICOM/NEMA website. The source
DOCX is never copied into the package data directory.
"""

from __future__ import annotations

import hashlib
import json
import re
import tempfile
import urllib.request
import zipfile
from pathlib import Path
from typing import Any

from ..http_utils import build_no_redirect_opener, validate_http_url

OFFICIAL_SOURCE_URL = "https://dicom.nema.org/medical/dicom/current/output/docx/part15.docx"

ATTRIBUTE_COLUMNS = [
    "name",
    "tag",
    "retired",
    "standard_composite_iod",
    "basic_profile",
    "retain_safe_private",
    "retain_uids",
    "retain_device_identity",
    "retain_institution_identity",
    "retain_patient_characteristics",
    "retain_longitudinal_full_dates",
    "retain_longitudinal_modified_dates",
    "clean_descriptors",
    "clean_structured_content",
    "clean_graphics",
]

CODE_COLUMNS = [
    "code_meaning",
    "code_value",
    "coding_scheme_designator",
    "value_type",
    "retired",
    "standard_template",
    "basic_profile",
    "retain_uids",
    "retain_device_identity",
    "retain_institution_identity",
    "retain_patient_characteristics",
    "retain_longitudinal_full_dates",
    "retain_longitudinal_modified_dates",
    "clean_descriptors",
]

DICOM_RIGHTS_NOTICE = (
    "DICOM® is the registered trademark of the National Electrical Manufacturers "
    "Association for its standards publications relating to digital communications "
    "of medical information, all rights reserved."
)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _text(cell: Any) -> str:
    return " ".join(cell.text.replace("\xa0", " ").replace("\u200b", "").split())


def _exact_tag(value: str) -> str | None:
    match = re.fullmatch(r"\(([0-9A-Fa-f]{4}),([0-9A-Fa-f]{4})\)", value.strip())
    return "".join(match.groups()).upper() if match else None


def _table(document: Any, first_header: str):
    for table in document.tables:
        if table.rows and _text(table.rows[0].cells[0]) == first_header:
            return table
    raise RuntimeError(f"Could not locate table beginning with {first_header!r}")


def _rows(table: Any, columns: list[str], *, attributes: bool) -> list[dict[str, str | None]]:
    output: list[dict[str, str | None]] = []
    for row in table.rows[1:]:
        values = [_text(cell) for cell in row.cells]
        if not any(values):
            continue
        if len(values) != len(columns):
            raise RuntimeError(f"Unexpected column count: {len(values)} instead of {len(columns)}")
        item: dict[str, str | None] = dict(zip(columns, values, strict=True))
        if attributes:
            item["tag_hex"] = _exact_tag(str(item["tag"]))
        output.append(item)
    return output


def _load_document(source: Path):
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - exercised by installation guidance
        raise RuntimeError(
            "PS3.15 table generation requires python-docx. Install the standards extra: "
            "pip install 'dicom-privacy-auditor[standards]'"
        ) from exc
    return Document(str(source))


def generate_tables(
    source: str | Path,
    *,
    edition: str,
    output: str | Path,
    source_url: str | None = None,
    minimum_attribute_rows: int = 600,
    minimum_code_rows: int = 200,
) -> tuple[Path, Path]:
    """Parse an official Part 15 DOCX into a user-local JSON rule cache."""

    source_path = Path(source).expanduser().resolve()
    output_path = Path(output).expanduser().resolve()
    if not source_path.is_file():
        raise FileNotFoundError(f"Source DOCX does not exist: {source_path}")

    document = _load_document(source_path)
    document_text = "\n".join(paragraph.text for paragraph in document.paragraphs[:150])
    expected_marker = f"DICOM PS3.15 {edition}"
    if expected_marker not in document_text:
        raise RuntimeError(
            f"Expected {expected_marker!r} was not found; refusing to generate mislabeled data"
        )

    attributes = _rows(_table(document, "Attribute Name"), ATTRIBUTE_COLUMNS, attributes=True)
    codes = _rows(_table(document, "Code Meaning"), CODE_COLUMNS, attributes=False)
    if len(attributes) < minimum_attribute_rows or len(codes) < minimum_code_rows:
        raise RuntimeError(f"Unexpected table sizes: E.1-1={len(attributes)}, E.1-2={len(codes)}")

    output_path.mkdir(parents=True, exist_ok=True)
    common = {
        "schema_version": "1.1",
        "standard": "DICOM PS3.15",
        "edition": edition,
        "source_url": source_url,
        "source_sha256": _sha256(source_path),
        "generated_locally": True,
        "redistributed_by_project": False,
        "rights_notice": DICOM_RIGHTS_NOTICE,
        "notice": (
            "This file was generated locally by the user from an official DICOM Standard "
            "document. It is not distributed as part of DICOM Privacy Auditor. Review NEMA/DICOM "
            "terms before redistributing this derived file."
        ),
    }
    attribute_path = output_path / f"ps315_{edition}_table_e1_1.json"
    code_path = output_path / f"ps315_{edition}_table_e1_2.json"
    payloads = [
        (
            attribute_path,
            {
                **common,
                "table": "E.1-1",
                "title": "Application Level Confidentiality Profile Attributes",
                "row_count": len(attributes),
                "columns": ATTRIBUTE_COLUMNS,
                "rules": attributes,
            },
        ),
        (
            code_path,
            {
                **common,
                "table": "E.1-2",
                "title": "Basic Application Confidentiality Profile Code Sequence Attributes",
                "row_count": len(codes),
                "columns": CODE_COLUMNS,
                "rules": codes,
            },
        ),
    ]
    for path, payload in payloads:
        temporary = path.with_suffix(path.suffix + ".tmp")
        temporary.write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
        temporary.replace(path)
    return attribute_path, code_path


def download_official_docx(
    url: str = OFFICIAL_SOURCE_URL,
    *,
    destination: str | Path,
    max_bytes: int = 100 * 1024 * 1024,
) -> Path:
    """Download an HTTPS DOCX without redirects and with a strict size limit."""

    validate_http_url(url, require_https=True)
    if max_bytes <= 0:
        raise ValueError("max_bytes must be positive")
    destination_path = Path(destination)
    destination_path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path = destination_path.with_suffix(destination_path.suffix + ".part")
    temporary_path.unlink(missing_ok=True)
    request = urllib.request.Request(
        url,
        headers={"User-Agent": "DICOM-Privacy-Auditor standards-cache-generator"},
    )
    opener = build_no_redirect_opener()
    written = 0
    try:
        with opener.open(request, timeout=60) as response, temporary_path.open("wb") as output:
            declared = int(response.headers.get("Content-Length", "0") or 0)
            if declared > max_bytes:
                raise RuntimeError("Standards source exceeded max_bytes")
            while chunk := response.read(min(1024 * 1024, max_bytes - written + 1)):
                written += len(chunk)
                if written > max_bytes:
                    raise RuntimeError("Standards source exceeded max_bytes")
                output.write(chunk)
        if not zipfile.is_zipfile(temporary_path):
            raise RuntimeError("Downloaded standards source is not a valid DOCX/ZIP file")
        temporary_path.replace(destination_path)
    except Exception:
        temporary_path.unlink(missing_ok=True)
        raise
    return destination_path


def download_and_generate(
    *,
    edition: str,
    output: str | Path,
    url: str = OFFICIAL_SOURCE_URL,
) -> tuple[Path, Path]:
    """Temporarily download the official DOCX, generate local tables, then delete the DOCX."""

    with tempfile.TemporaryDirectory(prefix="dicom-ps315-") as temporary:
        source = download_official_docx(url, destination=Path(temporary) / "part15.docx")
        return generate_tables(source, edition=edition, output=output, source_url=url)
